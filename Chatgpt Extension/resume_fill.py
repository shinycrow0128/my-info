import argparse
import copy
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

OUTPUT_FILENAME = "Final_Resume.docx"
PLACEHOLDER_RE = re.compile(r"<<\s*([^<>]+?)\s*>>")
BOLD_RE = re.compile(r"<b>(.*?)</b>", re.I | re.S)
VERIFY_MARKER = "[[VERIFY]]"


def script_folder() -> Path:
    return Path(__file__).resolve().parent


def clean_key(key: str) -> str:
    """Accept both 'Profile-Title' and '<<Profile-Title>>' JSON keys."""
    key = str(key).strip()
    match = re.fullmatch(r"<<\s*([^<>]+?)\s*>>", key)
    return match.group(1).strip() if match else key


def load_json(path: Path) -> dict[str, str]:
    """Load the flat JSON object expected by the resume prompt."""
    with path.open("r", encoding="utf-8-sig") as f:
        raw = json.load(f)

    if not isinstance(raw, dict):
        raise ValueError("JSON root must be one object.")

    data = {}
    for key, value in raw.items():
        if not isinstance(value, str):
            raise ValueError(f"JSON value for {key!r} must be a string.")
        data[clean_key(key)] = value
    return data


def find_one_json(folder: Path) -> Path:
    files = [p for p in folder.glob("*.json") if p.is_file()]
    if len(files) != 1:
        names = ", ".join(p.name for p in files) or "none"
        raise RuntimeError(f"Keep exactly one .json file beside this script. Found: {names}")
    return files[0]


def looks_like_template(path: Path) -> bool:
    try:
        doc = Document(path)
    except Exception:
        return False
    text = "\n".join(p.text for p in doc.paragraphs)
    normalized = re.sub(r"\s+", "", text)
    return "<<Profile-Title>>" in normalized and "<<Profile-Summary>>" in normalized


def find_template(folder: Path) -> Path:
    candidates = [
        p for p in folder.glob("*.docx")
        if p.is_file()
        and p.name.lower() != OUTPUT_FILENAME.lower()
        and not p.name.startswith("~$")
    ]
    matches = [p for p in candidates if looks_like_template(p)]
    if len(matches) != 1:
        names = ", ".join(p.name for p in matches or candidates) or "none"
        raise RuntimeError(f"Could not identify exactly one resume template. Found: {names}")
    return matches[0]


def base_run_properties(paragraph: Paragraph):
    """Copy the first run formatting so replaced text keeps the template font/style."""
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        return copy.deepcopy(paragraph.runs[0]._r.rPr)
    return None


def clear_paragraph(paragraph: Paragraph):
    """Remove text/runs but keep paragraph formatting, bullets, spacing, and tabs."""
    p = paragraph._p
    for child in list(p):
        if not child.tag.endswith("}pPr"):
            p.remove(child)


def add_run_with_base(paragraph: Paragraph, text: str, rpr, bold=False):
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, copy.deepcopy(rpr))
    if bold:
        run.bold = True


def write_markup(paragraph: Paragraph, text: str, rpr=None):
    """Write text and support only <b>...</b> inline bold tags."""
    if rpr is None:
        rpr = base_run_properties(paragraph)
    clear_paragraph(paragraph)

    pos = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > pos:
            add_run_with_base(paragraph, text[pos:match.start()], rpr)
        add_run_with_base(paragraph, match.group(1), rpr, bold=True)
        pos = match.end()
    if pos < len(text):
        add_run_with_base(paragraph, text[pos:], rpr)
    if not text:
        add_run_with_base(paragraph, "", rpr)


def remove_paragraph(paragraph: Paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def paragraph_from_xml(element, source: Paragraph) -> Paragraph:
    return Paragraph(element, source._parent)


def numbered_keys(data: dict[str, str], pattern: str) -> list[int]:
    regex = re.compile(pattern)
    values = []
    for key, value in data.items():
        match = regex.fullmatch(key)
        if match and value.strip():
            values.append(int(match.group(1)))
    return sorted(values)


def require_contiguous(numbers: list[int], label: str):
    if numbers and numbers != list(range(1, max(numbers) + 1)):
        raise ValueError(f"{label} numbering must be contiguous from 1. Found: {numbers}")


def fill_skills(doc: Document, data: dict[str, str]):
    """Fill the 3–4 skill rows already present in Cody/Russell templates."""
    indices = sorted({
        int(m.group(1))
        for key in data
        if (m := re.fullmatch(r"(?:Category|Skills)(\d+)", key))
    })
    require_contiguous(indices, "Skill category")
    if len(indices) > 4:
        raise ValueError("These templates support a maximum of 4 skill categories.")

    rows = []
    for i in indices:
        category = data.get(f"Category{i}", "").strip()
        skills = data.get(f"Skills{i}", "").strip()
        if bool(category) != bool(skills):
            raise ValueError(f"Category{i} and Skills{i} must both have values.")
        if category:
            rows.append((i, category, skills))

    template_rows = [
        p for p in doc.paragraphs
        if re.search(r"<<\s*Category\d+\s*>>", p.text)
    ]
    if rows and not template_rows:
        raise ValueError("No skill placeholders were found in the template.")

    for pos, (_, category, skills) in enumerate(rows):
        if pos >= len(template_rows):
            raise ValueError("Not enough skill rows in the template.")
        category = re.sub(r"</?b>", "", category, flags=re.I)
        write_markup(template_rows[pos], f"<b>{category}</b>: {skills}")

    for p in template_rows[len(rows):]:
        remove_paragraph(p)


def company_indices(data: dict[str, str]) -> list[int]:
    found = set()
    for key in data:
        m = re.fullmatch(r"Company(\d+)-(?:Subtitle|WorkSummary|Bullet\d+)", key)
        if m:
            found.add(int(m.group(1)))
    return sorted(found)


def fill_company_bullets(doc: Document, data: dict[str, str], company: int):
    """Fill existing 4 bullet rows and clone the last row for Bullet5+ when needed."""
    numbers = numbered_keys(data, rf"Company{company}-Bullet(\d+)")
    require_contiguous(numbers, f"Company{company} bullet")

    placeholders = [
        p for p in doc.paragraphs
        if re.fullmatch(rf"<<\s*Company{company}-Bullet\d+\s*>>", p.text.strip())
    ]
    if numbers and not placeholders:
        raise ValueError(f"No bullet placeholders found for Company{company}.")
    if not placeholders:
        return

    pristine_xml = copy.deepcopy(placeholders[-1]._p)
    pristine_rpr = base_run_properties(placeholders[-1])
    last_used = None

    for pos, number in enumerate(numbers):
        text = data[f"Company{company}-Bullet{number}"]
        if pos < len(placeholders):
            target = placeholders[pos]
            write_markup(target, text)
        else:
            new_xml = copy.deepcopy(pristine_xml)
            if last_used is None:
                raise RuntimeError("Cannot determine bullet insertion point.")
            last_used._p.addnext(new_xml)
            target = paragraph_from_xml(new_xml, placeholders[-1])
            write_markup(target, text, pristine_rpr)
        last_used = target

    for p in placeholders[len(numbers):]:
        remove_paragraph(p)


def replace_other_placeholders(doc: Document, data: dict[str, str]):
    """Replace title, summary, company subtitle/work summary, and any other remaining placeholders."""
    missing = set()
    for paragraph in doc.paragraphs:
        if "<<" not in paragraph.text:
            continue

        original = paragraph.text

        def repl(match):
            key = match.group(1).strip()
            if key not in data:
                missing.add(key)
                return match.group(0)
            return data[key]

        write_markup(paragraph, PLACEHOLDER_RE.sub(repl, original))

    if missing:
        raise ValueError("Missing JSON values for placeholders: " + ", ".join(sorted(missing)))


def warn_about_draft_metrics(data: dict[str, str]):
    keys = [key for key, value in data.items() if VERIFY_MARKER in value]
    if keys:
        print("\nWARNING: draft benchmark metrics found.")
        print("Confirm, replace, or remove every [[VERIFY]] item before using the resume as factual history.")
        print("Fields: " + ", ".join(keys))


def fill_resume(template_path: Path, json_path: Path, output_path: Path):
    data = load_json(json_path)
    warn_about_draft_metrics(data)

    doc = Document(template_path)
    fill_skills(doc, data)

    for company in company_indices(data):
        fill_company_bullets(doc, data, company)

    replace_other_placeholders(doc, data)
    doc.save(output_path)


def parse_args(argv=None):
    parser = argparse.ArgumentParser(
        description=(
            "Fill a resume template DOCX from the generator's resume_content.json. "
            "With no arguments it scans this script's own folder for one JSON and "
            "one template, which is the hand-run mode; the tracker server passes "
            "explicit paths so that several bids can be filled at the same time "
            "without sharing a folder."
        )
    )
    parser.add_argument("--json", dest="json_path", help="Path to the resume content JSON.")
    parser.add_argument("--template", dest="template_path", help="Path to the Temp_<name>.docx template.")
    parser.add_argument("--out", dest="output_path", help="Path of the DOCX to write.")
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    folder = script_folder()

    json_path = Path(args.json_path) if args.json_path else find_one_json(folder)
    template_path = Path(args.template_path) if args.template_path else find_template(folder)
    output_path = Path(args.output_path) if args.output_path else folder / OUTPUT_FILENAME

    for label, path in (("JSON", json_path), ("Template", template_path)):
        if not path.is_file():
            raise RuntimeError(f"{label} file not found: {path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"JSON:     {json_path.name}")
    print(f"Template: {template_path.name}")
    print(f"Output:   {output_path.name}")

    fill_resume(template_path, json_path, output_path)
    print(f"\nCreated: {output_path}")


if __name__ == "__main__":
    try:
        main()
    except (ValueError, RuntimeError, json.JSONDecodeError) as error:
        # Content problems - a missing placeholder value, a malformed JSON - are
        # the caller's to fix, so they read as one line instead of a traceback.
        print(f"ERROR: {error}", file=sys.stderr)
        sys.exit(1)
